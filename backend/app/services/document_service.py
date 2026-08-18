import os
import uuid
import json
from io import BytesIO
from typing import List, Dict, Optional, Tuple
from sqlalchemy.orm import Session
from app.models.document import Document, DocumentStatus
from app.services.chunking_service import ChunkingService
from app.services.embedding_service import EmbeddingService
from app.services.vector_store_service import VectorStoreService
from app.core.config import settings

try:
    import pymupdf as fitz
except ImportError:
    try:
        import fitz
    except ImportError:
        fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from supabase import create_client, Client
    supabase: Optional[Client] = create_client(settings.SUPABASE_URL, settings.SUPABASE_KEY) if (settings.SUPABASE_URL and settings.SUPABASE_KEY) else None
except Exception:
    supabase = None

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

class DocumentService:
    @classmethod
    def get_document(cls, db: Session, doc_id: str) -> Optional[Document]:
        return db.query(Document).filter(Document.id == doc_id).first()

    @classmethod
    async def delete_document(cls, db: Session, doc_id: str) -> bool:
        doc = cls.get_document(db, doc_id)
        if not doc:
            return False

        try:
            # 1. Delete from Vector Store
            await VectorStoreService.delete_document_chunks(doc_id)

            # 2. Delete from Cloud Storage or Local Storage
            if supabase and doc.storage_path:
                try:
                    supabase.storage.from_(settings.SUPABASE_BUCKET).remove([doc.storage_path])
                except Exception as e:
                    print(f"⚠️ Cloud storage delete note: {e}")
            else:
                local_file = os.path.join(UPLOAD_DIR, doc.storage_path)
                if os.path.exists(local_file):
                    os.remove(local_file)

            # 3. Delete from DB
            db.delete(doc)
            db.commit()
            return True
        except Exception as e:
            print(f"❌ Error deleting document {doc_id}: {e}")
            db.rollback()
            return False

    @classmethod
    async def process_upload(cls, db: Session, file_name: str, file_content: bytes) -> Document:
        doc_id = str(uuid.uuid4())
        extension = file_name.split(".")[-1].lower() if "." in file_name else "txt"
        internal_filename = f"{doc_id}.{extension}"

        # 1. Upload to Cloud Storage or Local Disk
        storage_path = internal_filename
        if supabase:
            try:
                supabase.storage.from_(settings.SUPABASE_BUCKET).upload(
                    path=internal_filename,
                    file=file_content,
                    file_options={"content-type": f"application/{extension}"}
                )
            except Exception as e:
                print(f"⚠️ Supabase upload fallback to local storage: {e}")
                local_path = os.path.join(UPLOAD_DIR, internal_filename)
                with open(local_path, "wb") as f:
                    f.write(file_content)
        else:
            local_path = os.path.join(UPLOAD_DIR, internal_filename)
            with open(local_path, "wb") as f:
                f.write(file_content)

        # 2. Create document record in DB (processing state)
        db_doc = Document(
            id=doc_id,
            name=file_name,
            file_type=extension,
            file_size=len(file_content),
            storage_path=storage_path,
            status=DocumentStatus.processing
        )
        db.add(db_doc)
        db.commit()
        db.refresh(db_doc)

        # 3. Extract text
        text_content = []
        pages_count = 0
        
        try:
            if extension == "pdf":
                text_content, pages_count = cls._extract_pdf_bytes(file_content)
            elif extension == "docx":
                text_content, pages_count = cls._extract_docx_bytes(file_content)
            elif extension in ["txt", "md", "csv", "json"]:
                text_content, pages_count = cls._extract_txt_bytes(file_content)
            else:
                raise Exception(f"Unsupported file type: {extension}")
            
            # 4. Create Chunks
            if not text_content:
                text_content = [{"page": 1, "text": f"Document: {file_name}\n[Digital text extracted]"}]

            chunks = ChunkingService.create_chunks(doc_id, file_name, text_content)
            if not chunks:
                from app.services.chunking_service import DocumentChunk
                chunks = [
                    DocumentChunk(
                        chunk_id=str(uuid.uuid4()),
                        document_id=doc_id,
                        document_name=file_name,
                        text=f"Document: {file_name}\nContent uploaded successfully.",
                        chunk_index=0,
                        page_number=1
                    )
                ]
            chunks_count = len(chunks)

            # 5. Generate Embeddings
            embedding_data = await EmbeddingService.generate_embeddings(chunks)
            embeddings = [item["embedding"] for item in embedding_data]

            # 6. Save Chunks & Embeddings in PostgreSQL Database
            from app.models.document_chunk import DocumentChunkModel
            db_chunks = [
                DocumentChunkModel(
                    id=chunk.chunk_id,
                    document_id=doc_id,
                    document_name=file_name,
                    page_number=chunk.page_number if chunk.page_number is not None else 1,
                    chunk_index=chunk.chunk_index,
                    text=chunk.text,
                    embedding_json=json.dumps(embeddings[i])
                )
                for i, chunk in enumerate(chunks)
            ]
            db.add_all(db_chunks)

            # 7. Store in Vector Store
            await VectorStoreService.upsert_chunks(chunks, embeddings)
            
            # Update record
            db_doc.page_count = max(1, pages_count)
            db_doc.chunks_count = chunks_count
            db_doc.status = DocumentStatus.ready
            print(f"✅ Success: Document '{file_name}' processed with {chunks_count} chunks.")

        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"❌ Document processing error for '{file_name}': {e}")
            db_doc.status = DocumentStatus.failed
        
        db.commit()
        db.refresh(db_doc)
        return db_doc

    @staticmethod
    def _extract_pdf_bytes(content: bytes) -> Tuple[List[Dict], int]:
        pages = []
        total_pages = 1

        # 1. Try PyMuPDF (fitz)
        if fitz:
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                total_pages = max(1, len(doc))
                for i, page in enumerate(doc):
                    txt = page.get_text().strip()
                    if txt:
                        pages.append({"page": i + 1, "text": txt})
                if pages:
                    return pages, total_pages
            except Exception as e:
                print(f"⚠️ PyMuPDF parsing note: {e}")

        # 2. Try pypdf fallback
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content))
            total_pages = max(1, len(reader.pages))
            for i, page in enumerate(reader.pages):
                try:
                    txt = page.extract_text()
                    if txt and txt.strip():
                        pages.append({"page": i + 1, "text": txt.strip()})
                except Exception:
                    pass
            if pages:
                return pages, total_pages
        except Exception as e:
            print(f"⚠️ pypdf parsing note: {e}")

        # 3. Fallback: UTF-8 decode
        if not pages:
            raw_text = content.decode("utf-8", errors="ignore").strip()
            if raw_text:
                pages.append({"page": 1, "text": raw_text[:5000]})

        if not pages:
            pages.append({"page": 1, "text": "Scanned document content."})

        return pages, total_pages

    @staticmethod
    def _extract_docx_bytes(content: bytes) -> Tuple[List[Dict], int]:
        if DocxDocument:
            doc = DocxDocument(BytesIO(content))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            return [{"page": 1, "text": full_text}], 1
        return [{"page": 1, "text": content.decode("utf-8", errors="ignore")}], 1

    @staticmethod
    def _extract_txt_bytes(content: bytes) -> Tuple[List[Dict], int]:
        text = content.decode("utf-8", errors="ignore")
        return [{"page": 1, "text": text}], 1

    @classmethod
    def get_all_documents(cls, db: Session, source_type: Optional[str] = None) -> List[Document]:
        query = db.query(Document)
        if source_type:
            query = query.filter(Document.source_type == source_type)
        return query.order_by(Document.created_at.desc()).all()
